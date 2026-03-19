"""Template generation helpers.

Provides functions to generate DOCX and PDF template files with mock content.
These functions are designed to be reused later to insert real AI analysis content.
"""
from datetime import datetime
import os
import io
try:
    from docx import Document
except Exception:
    Document = None

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
except Exception:
    canvas = None
    letter = (612, 792)

def build_mock_content(template_id, context=None):
    """Return a string with mock contract content for the given template_id.

    `context` can be a dict with keys to interpolate into the template later.
    """
    now = datetime.utcnow().strftime('%Y-%m-%d')
    ctx = context or {}
    party_a = ctx.get('party_a', 'Bên A - Công ty ABC')
    party_b = ctx.get('party_b', 'Bên B - Công ty XYZ')

    content = []
    content.append(f"Mẫu hợp đồng ({template_id.upper()})")
    content.append(f"Ngày: {now}")
    content.append("")
    content.append("Các bên:")
    content.append(f"1. {party_a}")
    content.append(f"2. {party_b}")
    content.append("")
    content.append("Điều 1. Đối tượng hợp đồng")
    content.append("Mô tả ngắn gọn về đối tượng hợp đồng.")
    content.append("")
    content.append("Điều 2. Giá và phương thức thanh toán")
    content.append("Các điều khoản về giá, cách thức thanh toán và lịch thanh toán.")
    content.append("")
    content.append("Điều 3. Thời hạn và chấm dứt")
    content.append("Thời hạn hợp đồng và điều kiện chấm dứt.")
    content.append("")
    content.append("Điều 4. Cam kết và bảo hành")
    content.append("Các cam kết của các bên và điều khoản bảo hành (nếu có).")
    content.append("")
    content.append("Ký tên:")
    content.append("Bên A: ______________________")
    content.append("Bên B: ______________________")

    return "\n".join(content)


def save_docx(content_text, output_path):
    """Save the given text content into a .docx file at output_path."""
    if Document is None:
        raise RuntimeError('python-docx not installed')

    doc = Document()
    lines = content_text.split('\n')
    # Use first line as title
    if lines:
        doc.add_heading(lines[0], level=1)
        for line in lines[1:]:
            doc.add_paragraph(line)

    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)


def save_pdf(content_text, output_path):
    """Save the given text content into a simple PDF at output_path."""
    if canvas is None:
        raise RuntimeError('reportlab not installed')

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    c = canvas.Canvas(output_path, pagesize=letter)
    width, height = letter
    y = height - 72
    lines = content_text.split('\n')
    for line in lines:
        c.drawString(72, y, line)
        y -= 14
        if y < 72:
            c.showPage()
            y = height - 72
    c.save()


def generate_template_file(template_id, output_path, context=None):
    """Generate a template file (docx/pdf) determined by output_path extension.

    Returns the output_path on success.
    """
    ext = output_path.rsplit('.', 1)[-1].lower()
    content = build_mock_content(template_id, context=context)

    if ext == 'docx':
        save_docx(content, output_path)
    elif ext == 'pdf':
        save_pdf(content, output_path)
    else:
        raise ValueError('Unsupported extension: ' + ext)

    return output_path
